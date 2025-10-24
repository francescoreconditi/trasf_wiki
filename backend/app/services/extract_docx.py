"""DOCX extraction service using python-docx.

This module handles extraction of text and images from Word documents.
"""

import zipfile
from pathlib import Path

from docx import Document

from app.models.dto import ExtractedData
from app.services.storage import save_image


def _extract_formatted_text(paragraph, image_map: dict) -> str:
    """Extract text from paragraph preserving bold and italic formatting.

    Args:
        paragraph: python-docx Paragraph object
        image_map: Dictionary mapping rId to image filename

    Returns:
        Text with formatting markers (BOLD:text, ITALIC:text, IMAGE:filename)
    """
    # Track images already processed to avoid duplicates
    images_found = set()

    # Collect runs with their formatting
    run_data = []
    for run in paragraph.runs:
        # Check if this run contains an image
        has_image = False
        try:
            if run._element.xpath(".//a:blip"):
                blips = run._element.xpath(".//a:blip")
                for blip in blips:
                    embed = blip.get(
                        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                    )
                    if embed and embed in image_map:
                        img_filename = image_map[embed]
                        if img_filename not in images_found:
                            run_data.append(("IMAGE", img_filename))
                            images_found.add(img_filename)
                        has_image = True
        except Exception:
            pass

        if has_image:
            continue

        text = run.text
        if not text:
            continue

        # Determine format type
        is_bold = run.bold
        is_italic = run.italic

        if is_bold and is_italic:
            format_type = "BOLDITALIC"
        elif is_bold:
            format_type = "BOLD"
        elif is_italic:
            format_type = "ITALIC"
        else:
            format_type = "NORMAL"

        run_data.append((format_type, text))

    # Merge consecutive runs with same formatting
    result_parts = []
    i = 0
    while i < len(run_data):
        format_type, content = run_data[i]

        if format_type == "IMAGE":
            result_parts.append(f"IMAGE:{content}")
            i += 1
            continue

        # Accumulate text for same format
        accumulated_text = [content]
        j = i + 1
        while j < len(run_data) and run_data[j][0] == format_type:
            accumulated_text.append(run_data[j][1])
            j += 1

        # Add formatted text
        full_text = "".join(accumulated_text)
        if format_type == "NORMAL":
            result_parts.append(full_text)
        else:
            result_parts.append(f"{format_type}:{full_text}")

        i = j

    # Check for paragraph-level images that weren't found in runs
    try:
        blips = paragraph._element.xpath(".//a:blip")
        for blip in blips:
            embed = blip.get(
                "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
            )
            if embed and embed in image_map:
                img_filename = image_map[embed]
                if img_filename not in images_found:
                    result_parts.append(f"IMAGE:{img_filename}")
                    images_found.add(img_filename)
    except Exception:
        pass

    return "".join(result_parts)


def extract_docx(file_path: Path, job_id: str | None = None) -> ExtractedData:
    """Extract text and images from DOCX file.

    Args:
        file_path: Path to DOCX file
        job_id: Optional job ID for organizing extracted images

    Returns:
        ExtractedData with text, images, and metadata
    """
    text_content = []
    images_list = []
    metadata = {}

    try:
        # Load document
        doc = Document(str(file_path))

        # Extract metadata (core properties)
        if hasattr(doc.core_properties, "title"):
            metadata["title"] = doc.core_properties.title or ""
        if hasattr(doc.core_properties, "author"):
            metadata["author"] = doc.core_properties.author or ""
        if hasattr(doc.core_properties, "subject"):
            metadata["subject"] = doc.core_properties.subject or ""

        # First, extract all images and create a map of rId -> filename
        image_map = _extract_images_and_create_map(file_path, doc, job_id)

        # Extract text from paragraphs with formatting
        for para in doc.paragraphs:
            # Check if paragraph has images even without text
            has_images_in_para = False
            try:
                blips = para._element.xpath(".//a:blip")
                if blips:
                    has_images_in_para = True
            except Exception:
                pass

            # Skip only if no text AND no images
            if not para.text.strip() and not has_images_in_para:
                continue

            # Detect heading styles (only if there's text)
            if para.text.strip() and para.style.name.startswith("Heading"):
                level = para.style.name.replace("Heading", "").strip()
                if level.isdigit():
                    text_content.append(f"HEADING{level}:{para.text.strip()}")
                else:
                    text_content.append(para.text.strip())
                continue

            # If paragraph has only images (no text), process it specially
            if not para.text.strip() and has_images_in_para:
                formatted_text = _extract_formatted_text(para, image_map)
                if formatted_text:
                    text_content.append(formatted_text)
                continue

            # Detect list items by checking numbering properties
            # Lists are identified by the presence of numPr in paragraph properties
            if para._p.pPr is not None and para._p.pPr.numPr is not None:
                num_pr = para._p.pPr.numPr
                ilvl = num_pr.ilvl  # Indentation level

                # Try to determine if it's numbered or bulleted
                # Numbered lists typically have numFmt, bulleted have bullets
                is_numbered = False
                try:
                    if hasattr(num_pr, "numId") and num_pr.numId is not None:
                        # Check the numbering definition
                        num_id_val = num_pr.numId.val
                        if num_id_val is not None:
                            # Try to get numbering format from document
                            # For simplicity, check if first char looks like number
                            text_preview = para.text.strip()[:5]
                            if text_preview and text_preview[0].isdigit():
                                is_numbered = True
                except Exception:
                    pass

                formatted_text = _extract_formatted_text(para, image_map)
                if is_numbered:
                    text_content.append(f"LIST_NUMBER:{formatted_text}")
                else:
                    text_content.append(f"LIST_BULLET:{formatted_text}")
                continue

            # Also check for "List" style names as fallback
            if "List" in para.style.name:
                formatted_text = _extract_formatted_text(para, image_map)
                # Default to bullet for List styles
                text_content.append(f"LIST_BULLET:{formatted_text}")
                continue

            # Regular paragraph - extract with formatting
            formatted_text = _extract_formatted_text(para, image_map)
            if formatted_text:
                text_content.append(formatted_text)

        # Extract text from tables
        for table in doc.tables:
            text_content.append("TABLE_START")
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells]
                text_content.append("ROW:" + " | ".join(row_cells))
            text_content.append("TABLE_END")

        # Images are now handled inline, so we keep an empty list
        # The image_map contains all images that were saved
        images_list = list(image_map.values())

    except Exception as e:
        raise ValueError(f"Failed to extract DOCX: {str(e)}")

    # Combine text
    full_text = "\n".join(text_content)

    return ExtractedData(text=full_text, images=images_list, metadata=metadata)


def _extract_images_and_create_map(
    file_path: Path, doc, job_id: str | None = None
) -> dict[str, str]:
    """Extract images from DOCX and create rId -> filename map.

    Args:
        file_path: Path to DOCX file
        doc: python-docx Document object
        job_id: Optional job ID for organizing images

    Returns:
        Dictionary mapping relationship IDs to saved image filenames
    """
    image_map = {}

    try:
        # Get relationships from document
        rels = doc.part.rels

        with zipfile.ZipFile(str(file_path), "r") as zip_ref:
            # DOCX stores images in word/media/
            media_files = [f for f in zip_ref.namelist() if f.startswith("word/media/")]

            print(f"[DOCX DEBUG] Found {len(media_files)} media files")
            print(f"[DOCX DEBUG] Total relationships: {len(rels)}")

            # Create a map of target -> file path
            target_to_file = {}
            for media_file in media_files:
                filename = media_file.split("/")[-1]
                target_to_file[filename] = media_file
                print(f"[DOCX DEBUG] Media file: {filename}")

            # Process relationships to map rId to saved image files
            image_rels = [r for r in rels.items() if "image" in r[1].reltype]
            print(f"[DOCX DEBUG] Found {len(image_rels)} image relationships")

            for rel_id, rel in rels.items():
                if "image" in rel.reltype:
                    print(f"[DOCX DEBUG] Processing image relationship: {rel_id}")
                    try:
                        # Get the target filename from the relationship
                        # rel.target_ref is a string like '../media/image1.png'
                        target_ref = (
                            rel.target_ref
                            if hasattr(rel, "target_ref")
                            else rel._target
                        )
                        if isinstance(target_ref, str):
                            target = target_ref.split("/")[-1]
                        else:
                            # If it's an object, try to get partname
                            target = (
                                target_ref.partname.split("/")[-1]
                                if hasattr(target_ref, "partname")
                                else str(target_ref).split("/")[-1]
                            )

                        if target in target_to_file:
                            print(f"[DOCX DEBUG] Found target {target} in media files")
                            # Read and save image
                            image_bytes = zip_ref.read(target_to_file[target])
                            ext = target.split(".")[-1].lower()
                            if ext not in ["png", "jpg", "jpeg", "gif", "bmp"]:
                                ext = "png"

                            # Save with descriptive name
                            base_name = f"docx_{target.rsplit('.', 1)[0]}"
                            image_path = save_image(image_bytes, ext, base_name, job_id)

                            # Keep the full URL path for frontend access
                            image_map[rel_id] = image_path
                            print(f"[DOCX DEBUG] Saved image: {image_path}")
                        else:
                            print(
                                f"[DOCX DEBUG] Target {target} NOT found in media files!"
                            )

                    except Exception as e:
                        print(
                            f"[DOCX ERROR] Error extracting image for rel {rel_id}: {e}"
                        )
                        import traceback

                        traceback.print_exc()

    except Exception as e:
        print(f"[DOCX ERROR] Error accessing DOCX images: {e}")
        import traceback

        traceback.print_exc()

    print(f"[DOCX DEBUG] Total images extracted: {len(image_map)}")
    return image_map
